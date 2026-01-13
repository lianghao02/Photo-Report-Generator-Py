from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import io
import os
import re
from utils import compress_image

# --- Layout Configuration ---
LAYOUT_STYLES = {
    'A4_Vertical': {
        'items_per_table': 1,
        'top_margin': 1.5,
        'bottom_margin': 1.9,
        'left_margin': 3.17,
        'right_margin': 3.17,
        'table_spacing_pt': 4,
        'table_spacing_font_pt': 1,
        'default_page_width': 21.0,
        'default_page_height': 29.7,
        'max_img_width': 14.4,
        'max_img_height': 9.8,
        'suffix_mode': False # Use [Key] without number
    },
    'A4_SideBySide': {
        'items_per_table': 2,
        'top_margin': 2.54,
        'bottom_margin': 2.54,
        'left_margin': 1.9,
        'right_margin': 1.9,
        'table_spacing_pt': 4,
        'table_spacing_font_pt': 1,
        'default_page_width': 21.0, # Assumed A4
        'default_page_height': 29.7,
        'max_img_width': 8.3,
        'max_img_height': 18.0,
        'suffix_mode': True # Use [Key 1], [Key 2]
    }
}

def analyze_docx_structure(template_path):
    try:
        doc = Document(template_path)
        output = [f"File: {os.path.basename(template_path)}"]
        output.append(f"Tables: {len(doc.tables)}")
        for t_idx, table in enumerate(doc.tables):
            output.append(f"\n[Table {t_idx}] {len(table.rows)} rows x {len(table.columns)} cols")
            for r_idx, row in enumerate(table.rows):
                row_txts = []
                for c_idx, cell in enumerate(row.cells):
                    txt = cell.text.strip().replace('\n', '\\n')
                    if txt:
                        row_txts.append(f"({r_idx},{c_idx}){txt}")
                if row_txts:
                     output.append(" | ".join(row_txts))
        return "\n".join(output)
    except Exception as e:
        return f"Error: {e}"

def create_photo_report(context, photos, template_path=None, layout_style='A4_Vertical'):
    if not template_path or not os.path.exists(template_path):
        return None
    
    # Load Config
    config = LAYOUT_STYLES.get(layout_style, LAYOUT_STYLES['A4_Vertical'])
    
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("模板中未發現表格")
        
    master_table = doc.tables[0]
    
    # Backup clean XML
    template_tbl_xml = deepcopy(master_table._element)
    
    # Clean up template: Remove everything AFTER the first table
    body_element = doc.element.body
    found_table = False
    elements_to_remove = []
    
    for element in body_element:
        if element == master_table._element:
            found_table = True
            continue
        if found_table:
            # Remove EVERYTHING, including sectPr. 
            # We will re-apply margins later, which creates a fresh sectPr.
            elements_to_remove.append(element)
            
    for el in elements_to_remove:
        body_element.remove(el)

    total_photos = len(photos)
    items_per_table = config['items_per_table']
    
    # Loop through photos in chunks
    for i in range(0, total_photos, items_per_table):
        # Determine table to use
        if i == 0:
            current_table = master_table
        else:
            # Append Spacer
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(config['table_spacing_pt'])
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Minimize font size
            run = p.add_run()
            run.font.size = Pt(config['table_spacing_font_pt'])
            
            # Clone Table
            new_tbl = deepcopy(template_tbl_xml)
            doc.element.body.append(new_tbl)
            current_table = doc.tables[-1]

        # Prepare Batch Data & Mapping
        # We need to constructing a single mapping for this table that includes all items in the batch
        # e.g. [圖片 1] -> photo[i], [圖片 2] -> photo[i+1]
        
        vals_for_slot = {}
        
        # 1. Global Context (Static strings)
        vals_for_slot[r"\[案由\]"] = context.get('案由', '')
        vals_for_slot[r"\[製作人\]"] = context.get('製作人', '')
        
        # For Side-by-Side (or multi-item) layout, [日期] often refers to the Global Header Date
        if config['items_per_table'] > 1:
             vals_for_slot[r"\[日期\]"] = context.get('日期', '')
        # If [日期] is meant to be the PHOTO date, we need to handle it per item.
        
        for idx, photo_data in enumerate(photos[i : i + items_per_table]):
            # Suffix: if items_per_table > 1, we use " 1", " 2". 
            # If items_per_table == 1, we use "" (empty).
            
            suffix = f" {idx+1}" if config['items_per_table'] > 1 else ""
            
            # Key Mapping
            # Photo Data Keys: 'date', 'time', 'location', 'no', 'desc', 'image'
            
            # [日期 1], [時間 1], ...
            vals_for_slot[fr"\[日期{suffix}\]"] = photo_data.get('date', '')
            vals_for_slot[fr"\[時間{suffix}\]"] = photo_data.get('time', '')
            vals_for_slot[fr"\[地點{suffix}\]"] = photo_data.get('location', '')
            vals_for_slot[fr"\[編號{suffix}\]"] = photo_data.get('no', '')
            vals_for_slot[fr"\[說明{suffix}\]"] = photo_data.get('desc', '')
            
            # Image is special object
            vals_for_slot[fr"\[圖片{suffix}\]"] = {'type': 'image', 'val': photo_data.get('image')}

        # Fill Data
        fill_slot(current_table, vals_for_slot, config)
    
    # --- Finalize Layout (Apply at the VERY END) ---
    # Ensure a section exists for the whole doc
    if doc.element.body.sectPr is None:
        doc.element.body.get_or_add_sectPr()

    # Apply Configured Margins
    for section in doc.sections:
        section.top_margin = Cm(config['top_margin'])
        section.bottom_margin = Cm(config['bottom_margin'])
        section.left_margin = Cm(config['left_margin'])
        section.right_margin = Cm(config['right_margin'])
        section.page_width = Cm(config['default_page_width'])
        section.page_height = Cm(config['default_page_height'])
        
        # --- HEADER ---
        header_text = context.get('header_text', '')
        if header_text:
            # Clear existing header content if any (optional, but cleaner)
            section.header.is_linked_to_previous = False
            # Remove existing paragraphs
            for p in section.header.paragraphs:
                p._element.getparent().remove(p._element)
                
            header_para = section.header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_para.add_run(header_text)
            set_run_font(run, '標楷體', 24)
            
        # --- FOOTER ---
        # Format: 第 X 頁 - 共 Y 頁
        section.footer.is_linked_to_previous = False
        for p in section.footer.paragraphs:
            p._element.getparent().remove(p._element)
            
        footer_para = section.footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Helper to add simple field
        def add_simple_field(paragraph, instr_text):
            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), instr_text)
            paragraph._element.append(fldSimple)

        # "第 "
        r1 = footer_para.add_run("第 ")
        set_run_font(r1, '標楷體', 12)
        
        # PAGE
        add_simple_field(footer_para, 'PAGE')
        
        # " 頁 - 共 "
        r2 = footer_para.add_run(" 頁 - 共 ")
        set_run_font(r2, '標楷體', 12)
        
        # NUMPAGES
        add_simple_field(footer_para, 'NUMPAGES')
        
        # " 頁"
        r3 = footer_para.add_run(" 頁")
        set_run_font(r3, '標楷體', 12)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def set_run_font(run, font_name='標楷體', size_pt=12):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def fill_slot(table, vals_for_slot, config):
    for row in table.rows:
        for cell in row.cells:
            for pattern, value in vals_for_slot.items():
                if re.search(pattern, cell.text, re.IGNORECASE):
                    
                    is_image = False
                    val_content = value
                    
                    if isinstance(value, dict) and value.get('type') == 'image':
                        is_image = True
                        val_content = value.get('val')
                        
                    if is_image:
                        cell.text = re.sub(pattern, "", cell.text, flags=re.IGNORECASE)
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        if val_content:
                            # Calculate Aspect Ratio and Dimensions
                            max_w = config['max_img_width']
                            max_h = config['max_img_height']
                            
                            img_w, img_h = val_content.size
                            img_ratio = img_w / img_h
                            target_ratio = max_w / max_h
                            
                            final_width = None
                            final_height = None
                            
                            if img_ratio > target_ratio:
                                # Width is the limiter
                                final_width = Cm(max_w)
                            else:
                                # Height is the limiter
                                final_height = Cm(max_h)
                            
                            img_stream = compress_image(val_content, max_size=(1600, 1600))
                            run.add_picture(img_stream, width=final_width, height=final_height)
                    else:
                        # Text Replace
                        # Note: We need to handle paragraphs carefully
                        for p in cell.paragraphs:
                            if re.search(pattern, p.text, re.IGNORECASE):
                                new_text = re.sub(pattern, str(val_content), p.text, flags=re.IGNORECASE)
                                p.text = new_text 
                                for run in p.runs:
                                    set_run_font(run, '標楷體', 12)
